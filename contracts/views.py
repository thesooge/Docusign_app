from django.shortcuts import render, redirect, get_object_or_404
from django.core.mail import send_mail
from django.urls import reverse
from docx import Document
from django.conf import settings
from django.http import HttpResponse
from django.views.generic import ListView
from django.contrib import messages
from django.utils import timezone
from datetime import timedelta
from urllib.parse import urlencode

import os
import base64
import requests
from docx2pdf import convert
import ctypes
import logging

from .models import Contract, DocusignProfile

logger = logging.getLogger(__name__)

if os.name == "nt":
    ole32 = ctypes.windll.ole32
    result = ole32.CoInitialize(None)
else:
    ole32 = None
    result = None

def docusign_login(request):
    redirect_uri = request.build_absolute_uri(reverse('docusign_callback'))
    print(redirect_uri)
    params = {
        "response_type": "code",
        "scope": "signature",
        "client_id": settings.DOCUSIGN_CLIENT_ID,
        "redirect_uri": redirect_uri
    }
    url = f"https://account-d.docusign.com/oauth/auth?{urlencode(params)}"
    return redirect(url)

def docusign_callback(request):
    code = request.GET.get("code")
    if not code:
        return HttpResponse("No code provided")

    redirect_uri = request.build_absolute_uri(reverse('docusign_callback'))
    data = {
        "grant_type": "authorization_code",
        "code": code,
        "client_id": settings.DOCUSIGN_CLIENT_ID,
        "client_secret": settings.DOCUSIGN_CLIENT_SECRET,
        "redirect_uri": redirect_uri
    }

    response = requests.post("https://account-d.docusign.com/oauth/token", data=data)

    if response.status_code == 200:
        token_data = response.json()
        access_token = token_data["access_token"]
        refresh_token = token_data["refresh_token"]
        expires_in = token_data["expires_in"]

        # Fetch account info from /userinfo endpoint
        userinfo_response = requests.get(
            "https://account-d.docusign.com/oauth/userinfo",
            headers={"Authorization": f"Bearer {access_token}"}
        )

        if userinfo_response.status_code != 200:
            return HttpResponse("Failed to get user info from DocuSign")

        userinfo = userinfo_response.json()
        account_info = userinfo["accounts"][0]
        account_id = account_info["account_id"]
        base_uri = account_info["base_uri"]

        # Save profile
        DocusignProfile.objects.update_or_create(
            user=request.user,
            defaults={
                "access_token": access_token,
                "refresh_token": refresh_token,
                "token_expiry": timezone.now() + timedelta(seconds=int(expires_in)),
                "account_id": account_id,
                "base_uri": base_uri
            }
        )

        return redirect("contract_instantiation")

    return HttpResponse("Failed to authenticate with DocuSign")

def get_user_token(user):
    profile = DocusignProfile.objects.filter(user=user).first()
    if not profile:
        return None

    if timezone.now() >= profile.token_expiry:
        data = {
            "grant_type": "refresh_token",
            "refresh_token": profile.refresh_token,
            "client_id": settings.DOCUSIGN_CLIENT_ID,
            "client_secret": settings.DOCUSIGN_CLIENT_SECRET
        }
        response = requests.post("https://account-d.docusign.com/oauth/token", data=data)
        if response.status_code == 200:
            token_data = response.json()
            profile.access_token = token_data["access_token"]
            profile.refresh_token = token_data["refresh_token"]
            profile.token_expiry = timezone.now() + timedelta(seconds=int(token_data["expires_in"]))
            profile.save()
        else:
            return None
    return profile.access_token, profile.account_id

def create_contract(request):
    if request.method == "POST":
        user_name = request.POST["user_name"]
        recipient_name = request.POST["recipient_name"]
        recipient_email = request.POST["recipient_email"]

        doc = Document()
        doc.add_heading("Contract Agreement", level=1)
        doc.add_paragraph(f"Party 1: {user_name}")
        doc.add_paragraph(f"Party 2: {recipient_name}")
        doc.add_paragraph("\nThis agreement is binding and requires signatures.")

        contract_filename = f"contract_{user_name}_{recipient_name}.docx"
        contract_path = os.path.join(settings.MEDIA_ROOT, contract_filename)
        doc.save(contract_path)
        return redirect(reverse("send_to_docusign") + f"?contract_path={contract_filename}&recipient_email={recipient_email}&user_name={user_name}&recipient_name={recipient_name}")

    return render(request, "contracts/contract_form.html")

def encode_file_to_base64(file_path):
    if not os.path.exists(file_path):
        return None
    with open(file_path, "rb") as f:
        return base64.b64encode(f.read()).decode("utf-8")

def submit_contract_to_docusign(request):
    user = request.user
    token_account = get_user_token(user)
    if not token_account:
        return redirect("docusign_login")

    access_token, account_id = token_account

    user_name = request.GET.get("user_name")
    recipient_name = request.GET.get("recipient_name")
    contract_filename = request.GET.get("contract_path")
    recipient_email = request.GET.get("recipient_email")

    if not all([contract_filename, recipient_email, user_name, recipient_name]):
        messages.error(request, "Missing required information.")
        return redirect("contract_instantiation")

    contract_path = os.path.join(settings.MEDIA_ROOT, contract_filename)
    pdf_path = contract_path.replace(".docx", ".pdf")

    try:
        convert(contract_path, pdf_path)
    except Exception as e:
        return HttpResponse(f"Conversion error: {str(e)}")

    encoded_pdf = encode_file_to_base64(pdf_path)
    if not encoded_pdf:
        return HttpResponse("Error reading PDF.")

    contract = Contract.objects.create(
        user_name=user_name,
        recipient_email=recipient_email,
        recipient_name=recipient_name,
        contract_file=contract_filename
    )

    envelope_data = {
        "emailSubject": "Contract Agreement - Please Sign",
        "documents": [{
            "documentBase64": encoded_pdf,
            "name": "Contract Agreement",
            "fileExtension": "pdf",
            "documentId": "1"
        }],
        "recipients": {
            "signers": [{
                "email": recipient_email,
                "name": "Recipient",
                "recipientId": "1",
                "tabs": {
                    "signHereTabs": [{"xPosition": "200", "yPosition": "500", "documentId": "1", "pageNumber": "1"}]
                }
            }]
        },
        "status": "sent"
    }

    url = f"https://demo.docusign.net/restapi/v2.1/accounts/{account_id}/envelopes"
    headers = {
        "Authorization": f"Bearer {access_token}",
        "Content-Type": "application/json"
    }

    response = requests.post(url, headers=headers, json=envelope_data)
    if response.status_code == 201:
        envelope_id = response.json().get("envelopeId")
        contract.document_id = envelope_id
        contract.save()
        notify_recipient(recipient_email, envelope_id)
        return redirect("success_page")
    return HttpResponse("Error sending contract: " + response.text)

def notify_recipient(email, contract_url):
    subject = "Contract Agreement - Please Sign"
    message = f"Please sign the contract using the following link: {contract_url}"
    try:
        send_mail(subject, message, settings.DEFAULT_FROM_EMAIL, [email])
    except Exception as e:
        logger.error(f"Error sending email to {email}: {str(e)}")

def success_page(request):
    return render(request, "contracts/success.html")

def is_contract_signed(contract):
    profile = DocusignProfile.objects.filter(user__username=contract.user_name).first()
    if not profile:
        return False
    token, account_id = get_user_token(profile.user)
    url = f"https://demo.docusign.net/restapi/v2.1/accounts/{account_id}/envelopes/{contract.document_id}"
    headers = {
        "Authorization": f"Bearer {token}",
        "Content-Type": "application/json"
    }
    response = requests.get(url, headers=headers)
    if response.status_code == 200:
        if response.json().get("status") == "completed":
            contract.is_signed = True
            contract.save()
            return True
    return False

class ContractListView(ListView):
    model = Contract
    template_name = "contracts/contract_list.html"

    def post(self, request, *args, **kwargs):
        contract_id = request.POST.get("contract_id")
        if contract_id:
            contract = get_object_or_404(Contract, id=contract_id)
            if is_contract_signed(contract):
                messages.success(request, f"Contract '{contract.id}' is signed.")
            else:
                messages.warning(request, f"Contract '{contract.id}' is not signed.")
        return self.get(request, *args, **kwargs)